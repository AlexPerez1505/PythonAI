import argparse
import json
import os
import re
import sys
import time
from io import BytesIO
from pathlib import Path
from typing import Dict, Any, List, Optional

import requests
from dotenv import load_dotenv
from pypdf import PdfReader, PdfWriter

load_dotenv()

AZURE_ENDPOINT = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT", "").rstrip("/")
AZURE_KEY = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY", "")
AZURE_API_VERSION = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_API_VERSION", "2024-11-30")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "https://api.openai.com").rstrip("/")


# ════════════════════════════════════════════════════════════
# HELPERS
# ════════════════════════════════════════════════════════════

def _log(message: str) -> None:
    print(message, file=sys.stderr, flush=True)


def _is_size_error(message: str) -> bool:
    text = (message or "").lower()
    return "invalidcontentlength" in text or "too large" in text


def _norm(v: Any) -> Optional[str]:
    if v is None:
        return None

    t = re.sub(r"\s+", " ", str(v)).strip()
    return t or None


def _is_empty_value(value: Any) -> bool:
    text = _norm(value) or ""
    low = text.lower()

    return (
        not text
        or low in [
            "no se encontro informacion",
            "no se encontró información",
            "no se encontro información",
            "no se encontró informacion",
            "no aplica",
            "n/a",
            "na",
            "-",
            "sin informacion",
            "sin información",
        ]
    )


# ════════════════════════════════════════════════════════════
# AZURE DOCUMENT INTELLIGENCE
# ════════════════════════════════════════════════════════════

def _analyze_single_pdf_bytes(file_bytes: bytes, model: str = "prebuilt-layout") -> Dict[str, Any]:
    if not AZURE_ENDPOINT:
        raise Exception("Falta AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT en .env")

    if not AZURE_KEY:
        raise Exception("Falta AZURE_DOCUMENT_INTELLIGENCE_KEY en .env")

    url = f"{AZURE_ENDPOINT}/documentintelligence/documentModels/{model}:analyze?api-version={AZURE_API_VERSION}"

    headers = {
        "Ocp-Apim-Subscription-Key": AZURE_KEY,
        "Content-Type": "application/pdf",
    }

    response = requests.post(url, headers=headers, data=file_bytes, timeout=180)

    if not response.ok:
        raise Exception(f"Azure analyze error: {response.status_code} - {response.text}")

    operation_location = response.headers.get("operation-location")

    if not operation_location:
        raise Exception("Azure no devolvio operation-location")

    while True:
        poll = requests.get(
            operation_location,
            headers={"Ocp-Apim-Subscription-Key": AZURE_KEY},
            timeout=180,
        )

        if not poll.ok:
            raise Exception(f"Azure poll error: {poll.status_code} - {poll.text}")

        data = poll.json()
        status = data.get("status")

        if status in ["notStarted", "running"]:
            time.sleep(2)
            continue

        if status != "succeeded":
            raise Exception(f"Azure status no exitoso: {data}")

        return data


def _split_pdf_bytes(file_bytes: bytes, pages_per_chunk: int) -> List[bytes]:
    reader = PdfReader(BytesIO(file_bytes))
    total_pages = len(reader.pages)

    chunks = []

    for start in range(0, total_pages, pages_per_chunk):
        end = min(start + pages_per_chunk, total_pages)
        writer = PdfWriter()

        for i in range(start, end):
            writer.add_page(reader.pages[i])

        output = BytesIO()
        writer.write(output)
        chunks.append(output.getvalue())

    return chunks


def _merge_azure_results(results: List[Dict[str, Any]]) -> Dict[str, Any]:
    merged_content = []
    merged_pages = []
    merged_tables = []
    page_offset = 0

    for r in results:
        ar = r.get("analyzeResult", {}) or {}

        if ar.get("content"):
            merged_content.append(ar["content"])

        pages = ar.get("pages", []) or []

        for page in pages:
            pcopy = dict(page)

            if "pageNumber" in pcopy:
                pcopy["pageNumber"] = page_offset + pcopy["pageNumber"]

            merged_pages.append(pcopy)

        for table in ar.get("tables", []) or []:
            merged_tables.append(table)

        page_offset += len(pages)

    return {
        "status": "succeeded",
        "analyzeResult": {
            "content": "\n\n".join(merged_content),
            "pages": merged_pages,
            "tables": merged_tables,
        },
    }


def analyze_pdf_with_auto_split(file_bytes: bytes, pages_per_chunk: int = 5) -> Dict[str, Any]:
    try:
        _log("Azure: intentando PDF completo...")
        return _analyze_single_pdf_bytes(file_bytes)
    except Exception as e:
        if not _is_size_error(str(e)):
            raise

        _log("PDF rechazado por tamano, dividiendo...")

    chunk_size = pages_per_chunk

    while chunk_size >= 1:
        _log(f"Probando con chunks de {chunk_size} paginas...")

        chunks = _split_pdf_bytes(file_bytes, chunk_size)
        partials = []
        failed_size = False

        for idx, ch_bytes in enumerate(chunks, start=1):
            try:
                _log(f"Procesando chunk {idx}/{len(chunks)}...")
                partials.append(_analyze_single_pdf_bytes(ch_bytes))
            except Exception as e:
                if _is_size_error(str(e)) and chunk_size > 1:
                    _log(f"Chunk {idx} aun grande, reduciendo...")
                    failed_size = True
                    break

                raise

        if not failed_size:
            return _merge_azure_results(partials)

        chunk_size = chunk_size // 2

    raise Exception("Azure rechazo el PDF incluso dividido al minimo.")


def azure_result_to_page_marked_text(analyze_result: Dict[str, Any]) -> str:
    pages = analyze_result.get("pages", []) or []
    page_blocks = []

    for page in pages:
        page_number = page.get("pageNumber")
        lines = page.get("lines", []) or []

        line_texts = []

        for line in lines:
            content = _norm(line.get("content"))
            if content:
                line_texts.append(content)

        if line_texts:
            page_blocks.append(
                f"[PAGINA {page_number}]\n" + "\n".join(line_texts)
            )

    if page_blocks:
        return "\n\n".join(page_blocks)

    return analyze_result.get("content", "") or ""


# ════════════════════════════════════════════════════════════
# DOCX
# ════════════════════════════════════════════════════════════

def extract_docx_text(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise Exception("Falta paquete 'python-docx'. Instala: pip install python-docx")

    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


# ════════════════════════════════════════════════════════════
# EVIDENCE / CITAS FALLBACK
# ════════════════════════════════════════════════════════════

def _extract_source_and_page(raw_text: str, pos: int) -> Dict[str, Any]:
    before = raw_text[:pos]

    doc_matches = list(re.finditer(r"--- DOCUMENTO:\s*(.*?)\s*---", before))
    page_matches = list(re.finditer(r"\[PAGINA\s+(\d+)\]", before, flags=re.IGNORECASE))

    fuente = doc_matches[-1].group(1).strip() if doc_matches else ""
    pagina = int(page_matches[-1].group(1)) if page_matches else None

    return {
        "fuente": fuente,
        "pagina": pagina,
    }


def _make_quote(raw_text: str, pos: int, length: int = 350) -> str:
    start = max(0, pos - 100)
    end = min(len(raw_text), pos + length)

    quote = raw_text[start:end]
    quote = re.sub(r"--- DOCUMENTO:\s*.*?\s*---", "", quote)
    quote = re.sub(r"\[PAGINA\s+\d+\]", "", quote, flags=re.IGNORECASE)
    quote = re.sub(r"\s+", " ", quote).strip()

    return quote[:350]


def _find_evidence_for_value(raw_text: str, value: Any) -> Optional[Dict[str, Any]]:
    value_text = _norm(value)

    if not value_text or _is_empty_value(value_text):
        return None

    raw_low = raw_text.lower()
    value_low = value_text.lower()

    pos = raw_low.find(value_low)

    if pos < 0:
        words = [
            w for w in re.findall(r"[\wáéíóúñü]+", value_low, flags=re.IGNORECASE)
            if len(w) >= 4
        ]

        for size in range(min(8, len(words)), 2, -1):
            for i in range(0, len(words) - size + 1):
                phrase_words = words[i:i + size]
                pattern = r"\b" + r"\W+".join(map(re.escape, phrase_words)) + r"\b"
                match = re.search(pattern, raw_low, flags=re.IGNORECASE)

                if match:
                    pos = match.start()
                    break

            if pos >= 0:
                break

    if pos < 0:
        return None

    meta = _extract_source_and_page(raw_text, pos)

    return {
        "cita": _make_quote(raw_text, pos),
        "fuente": meta["fuente"],
        "pagina": meta["pagina"],
    }


def ensure_ficha_resumen_citations(structured: Dict[str, Any], raw_text: str) -> Dict[str, Any]:
    if not isinstance(structured, dict):
        return structured

    citas = structured.get("citas")

    if not isinstance(citas, dict):
        citas = {}

    ficha = structured.get("ficha") or {}
    fechas = structured.get("fechas_clave") or {}
    resumen = structured.get("resumen_ejecutivo") or []

    if isinstance(ficha, dict):
        for field, value in ficha.items():
            key = f"ficha.{field}"

            if key not in citas and not _is_empty_value(value):
                evidence = _find_evidence_for_value(raw_text, value)

                if evidence:
                    citas[key] = evidence

    if isinstance(fechas, dict):
        for field, value in fechas.items():
            key = f"fechas_clave.{field}"

            if key not in citas and not _is_empty_value(value):
                evidence = _find_evidence_for_value(raw_text, value)

                if evidence:
                    citas[key] = evidence

    if isinstance(resumen, list):
        for idx, item in enumerate(resumen):
            if not isinstance(item, dict):
                continue

            key = f"resumen_ejecutivo.{idx}"
            answer = item.get("respuesta")

            if key not in citas and not _is_empty_value(answer):
                evidence = _find_evidence_for_value(raw_text, answer)

                if evidence:
                    citas[key] = evidence

    structured["citas"] = citas
    return structured


def ensure_checklist_citations(structured: Dict[str, Any], raw_text: str) -> Dict[str, Any]:
    if not isinstance(structured, dict):
        return structured

    checklist = structured.get("checklist_sugerido")

    if not isinstance(checklist, list):
        return structured

    for item in checklist:
        if not isinstance(item, dict):
            continue

        if item.get("fuente") and item.get("pagina") and item.get("cita"):
            continue

        search_value = (
            item.get("requisito")
            or item.get("descripcion")
            or item.get("formato")
        )

        evidence = _find_evidence_for_value(raw_text, search_value)

        if evidence:
            item.setdefault("cita", evidence.get("cita"))
            item.setdefault("fuente", evidence.get("fuente"))
            item.setdefault("pagina", evidence.get("pagina"))

    return structured



# ════════════════════════════════════════════════════════════
# DETERMINISTIC FALLBACKS / BUSQUEDA DIRECTA EN TEXTO AZURE
# ════════════════════════════════════════════════════════════

def _clean_extracted_value(value: Any, max_len: int = 400) -> Optional[str]:
    text = _norm(value)

    if not text:
        return None

    text = re.sub(r"\s+", " ", text).strip(" \t\n\r\0\x0B:-–—.;")

    if not text:
        return None

    return text[:max_len].strip()


def _first_regex(raw_text: str, patterns: List[str], max_len: int = 400) -> Optional[Dict[str, Any]]:
    if not raw_text:
        return None

    for pattern in patterns:
        match = re.search(pattern, raw_text, flags=re.IGNORECASE | re.DOTALL | re.UNICODE)

        if not match:
            continue

        value = match.group(1) if match.groups() else match.group(0)
        clean = _clean_extracted_value(value, max_len=max_len)

        if not clean or _is_empty_value(clean):
            continue

        pos = match.start(1) if match.groups() else match.start()
        meta = _extract_source_and_page(raw_text, pos)

        return {
            "value": clean,
            "cita": _make_quote(raw_text, pos),
            "fuente": meta.get("fuente", ""),
            "pagina": meta.get("pagina"),
        }

    return None


def _set_if_empty(target: Dict[str, Any], key: str, found: Optional[Dict[str, Any]], citas: Dict[str, Any]) -> None:
    if not found:
        return

    if _is_empty_value(target.get(key)):
        target[key] = found.get("value")

    cita_key = f"ficha.{key}"

    if cita_key not in citas and not _is_empty_value(found.get("value")):
        citas[cita_key] = {
            "cita": found.get("cita"),
            "fuente": found.get("fuente"),
            "pagina": found.get("pagina"),
        }


def _find_known_phrase(raw_text: str, phrases: List[str]) -> Optional[Dict[str, Any]]:
    raw_low = raw_text.lower()

    for phrase in phrases:
        pos = raw_low.find(phrase.lower())

        if pos < 0:
            continue

        meta = _extract_source_and_page(raw_text, pos)

        return {
            "value": phrase,
            "cita": _make_quote(raw_text, pos),
            "fuente": meta.get("fuente", ""),
            "pagina": meta.get("pagina"),
        }

    return None


def _extract_likely_organismo(raw_text: str) -> Optional[Dict[str, Any]]:
    known = [
        "Consejo Nacional de Fomento Educativo",
        "Secretaría de Educación Pública",
        "Secretaria de Educación Pública",
        "Secretaría de Marina",
        "Secretaria de Marina",
        "Instituto Mexicano del Seguro Social",
        "Instituto de Seguridad y Servicios Sociales de los Trabajadores del Estado",
    ]

    found = _find_known_phrase(raw_text, known)

    if found:
        if found["value"].lower() == "secretaria de educación pública":
            found["value"] = "Secretaría de Educación Pública"
        return found

    conafe_pos = raw_text.lower().find("conafe")

    if conafe_pos >= 0:
        meta = _extract_source_and_page(raw_text, conafe_pos)
        return {
            "value": "Consejo Nacional de Fomento Educativo",
            "cita": _make_quote(raw_text, conafe_pos),
            "fuente": meta.get("fuente", ""),
            "pagina": meta.get("pagina"),
        }

    return _first_regex(raw_text, [
        r"(?:convocante|dependencia|entidad|organismo)\s*[:\-–—]?\s*([A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚÜÑáéíóúüñ\s\.,]{8,160})",
    ], 180)


def _extract_payment_conditions(raw_text: str) -> Optional[Dict[str, Any]]:
    return _first_regex(raw_text, [
        r"(?:\d+\.\s*)?CONDICIONES\s+Y\s+FORMAS\s+DE\s+PAGO\s*(.+?)(?=\s+(?:\d+\.\s+[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{6,}|P[aá]gina\s+\d+|--- DOCUMENTO:|$))",
        r"(?:condiciones|forma|formas)\s+de\s+pago\s*[:\-–—]?\s*(.+?)(?=\s+(?:\d+\.\s+[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{6,}|P[aá]gina\s+\d+|--- DOCUMENTO:|$))",
        r"(el\s+pago\s+correspondiente\s+se\s+realizar[aá]\s+.+?)(?=\s+(?:\d+\.\s+[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{6,}|P[aá]gina\s+\d+|--- DOCUMENTO:|$))",
    ], 800)


def _extract_currency(raw_text: str) -> Optional[Dict[str, Any]]:
    found = _first_regex(raw_text, [
        r"moneda\s+nacional\s*\(([^\)]{3,80})\)",
        r"\b(pesos\s+mexicanos)\b",
        r"\b(moneda\s+nacional)\b",
    ], 120)

    if not found:
        return None

    value = found.get("value") or ""
    low = value.lower()

    if "pesos mexicanos" in low:
        found["value"] = "Pesos mexicanos"
    elif "moneda nacional" in low:
        found["value"] = "Moneda nacional"
    else:
        found["value"] = f"Moneda nacional ({value})"

    return found


def apply_deterministic_fallbacks(structured: Dict[str, Any], raw_text: str) -> Dict[str, Any]:
    """
    Rellena campos que la IA dejó vacíos usando búsqueda directa sobre el texto
    que Azure ya leyó con marcas de página. Esto evita perder datos evidentes
    como moneda, condiciones de pago, organismo u objeto.
    """
    if not isinstance(structured, dict):
        structured = {}

    ficha = structured.get("ficha")

    if not isinstance(ficha, dict):
        ficha = {}

    citas = structured.get("citas")

    if not isinstance(citas, dict):
        citas = {}

    raw_text = raw_text or ""

    tipo_evento = _find_known_phrase(raw_text, [
        "Invitación a cuando menos tres personas",
        "Licitación Pública Nacional",
        "Licitación Pública Internacional",
        "Licitación Pública",
        "Adjudicación Directa",
    ])

    numero = _first_regex(raw_text, [
        r"(?:n[uú]mero\s+de\s+licitaci[oó]n|no\.?\s*(?:de\s*)?(?:licitaci[oó]n|procedimiento)|procedimiento\s*(?:no\.?|n[uú]m\.?)|expediente)\s*[:#]?\s*([A-Z0-9][A-Z0-9\/\-.]{4,})",
        r"\b((?:IA|LA|LPN|LPI|AA|AD)[\-\/]?[A-Z0-9\-\/\.]{5,})\b",
    ], 90)

    organismo = _extract_likely_organismo(raw_text)

    objeto = _first_regex(raw_text, [
        r"objeto\s+(?:de\s+)?(?:la\s+)?(?:licitaci[oó]n|contrataci[oó]n|procedimiento)\s*[:\-–—\n]\s*([^.;]{12,320})",
        r"(adquisici[oó]n\s+de\s+(?:materiales|bienes|servicios|[úu]tiles)[^.;]{5,260})",
        r"(contrataci[oó]n\s+de\s+(?:servicios|bienes|materiales)[^.;]{5,260})",
        r"(servicio\s+de\s+[^.;]{10,260})",
    ], 320)

    medio = None

    if re.search(r"\b(electr[oó]nic[ao]|CompraNet)\b", raw_text, flags=re.IGNORECASE):
        medio = _find_known_phrase(raw_text, ["CompraNet", "electrónica", "electronica"])
        if medio:
            medio["value"] = "Electrónica"
    elif re.search(r"\bpresencial\b", raw_text, flags=re.IGNORECASE):
        medio = _find_known_phrase(raw_text, ["presencial"])
        if medio:
            medio["value"] = "Presencial"
    elif tipo_evento:
        medio = dict(tipo_evento)

    moneda = _extract_currency(raw_text)
    condiciones_pago = _extract_payment_conditions(raw_text)

    _set_if_empty(ficha, "numero_licitacion", numero, citas)
    _set_if_empty(ficha, "tipo_evento", tipo_evento, citas)
    _set_if_empty(ficha, "organismo", organismo, citas)
    _set_if_empty(ficha, "objeto_licitacion", objeto, citas)
    _set_if_empty(ficha, "objeto", objeto, citas)
    _set_if_empty(ficha, "medio_participacion", medio, citas)
    _set_if_empty(ficha, "moneda_pago", moneda, citas)
    _set_if_empty(ficha, "condiciones_pago", condiciones_pago, citas)

    structured["ficha"] = ficha
    structured["citas"] = citas
    return structured

def postprocess_structured_data(structured: Dict[str, Any], raw_text: str) -> Dict[str, Any]:
    structured = apply_deterministic_fallbacks(structured, raw_text)
    structured = ensure_ficha_resumen_citations(structured, raw_text)
    structured = ensure_checklist_citations(structured, raw_text)
    return structured


# ════════════════════════════════════════════════════════════
# OPENAI STRUCTURER
# ════════════════════════════════════════════════════════════

LICITACION_SYSTEM_PROMPT = (
    "Eres un asistente experto en extraccion de informacion de licitaciones publicas mexicanas. "
    "Debes responder UNICAMENTE JSON valido, sin markdown, sin explicaciones y sin texto adicional."
)


def _build_licitacion_prompt(raw_text: str) -> str:
    compact = (raw_text or "")[:140000]

    return f"""
Analiza EXCLUSIVAMENTE el texto de los documentos proporcionados y devuelve UN SOLO JSON válido.

ESTRUCTURA OBLIGATORIA:
{{
  "ficha": {{
    "numero_licitacion": "",
    "tipo_evento": "",
    "organismo": "",
    "objeto_licitacion": "",
    "medio_participacion": "",
    "moneda_pago": "",
    "condiciones_pago": ""
  }},
  "fechas_clave": {{
    "fecha_publicacion": "",
    "junta_aclaraciones": "",
    "presentacion_apertura": "",
    "fallo": "",
    "vigencia_contrato": ""
  }},
  "resumen_ejecutivo": [
    {{"pregunta": "¿Cuánto tiempo tengo para implementar?", "respuesta": ""}},
    {{"pregunta": "¿Es necesario demostrar experiencia previa o acreditar experiencia?", "respuesta": ""}},
    {{"pregunta": "¿Se mencionan penas convencionales, multas, deducciones u otras sanciones en caso de incumplimiento?", "respuesta": ""}},
    {{"pregunta": "¿Cuál es el periodo de garantía a ofertar?", "respuesta": ""}},
    {{"pregunta": "¿Cuál es el sistema de evaluación?", "respuesta": ""}},
    {{"pregunta": "¿Se requieren cartas de apoyo?", "respuesta": ""}},
    {{"pregunta": "¿Se deben entregar muestras físicas?", "respuesta": ""}},
    {{"pregunta": "¿Es necesario entregar documentación regulatoria?", "respuesta": ""}},
    {{"pregunta": "¿A qué hospitales o instituciones se deben entregar los productos o prestar los servicios?", "respuesta": ""}},
    {{"pregunta": "¿Existe subrogación en caso de fallas del equipo?", "respuesta": ""}},
    {{"pregunta": "¿Se requiere la documentación técnica en español o se permiten traducciones simples?", "respuesta": ""}},
    {{"pregunta": "¿Cómo se realiza la adjudicación?", "respuesta": ""}},
    {{"pregunta": "¿Se menciona si el evento está bajo tratados de libre comercio?", "respuesta": ""}},
    {{"pregunta": "¿Cuál es la vigencia o duración del contrato?", "respuesta": ""}},
    {{"pregunta": "¿Cuáles son los plazos de entrega y las condiciones para cumplir con las entregas?", "respuesta": ""}},
    {{"pregunta": "¿Es necesario realizar una visita a las instalaciones de la convocante?", "respuesta": ""}}
  ],
  "eventos": {{
    "comentarios": "",
    "vigencias": [
      {{"label": "", "risk": "NULO", "value": "", "fuente": "", "pagina": null, "cita": ""}}
    ],
    "plazos_ejecucion": [
      {{"label": "", "risk": "NULO", "value": "", "fuente": "", "pagina": null, "cita": ""}}
    ]
  }},
  "matriz": {{
    "comentarios": "",
    "secciones": [
      {{
        "title": "",
        "icons": ["yellow"],
        "items": [
          {{
            "question": "",
            "risk": "NULO",
            "answer": "",
            "justificacion": "",
            "citas": [""],
            "fuente": "",
            "pagina": null
          }}
        ]
      }}
    ]
  }},
  "financiero": {{
    "comentarios": "",
    "secciones": [
      {{
        "title": "",
        "icons": ["yellow"],
        "items": [
          {{
            "question": "",
            "risk": "NULO",
            "answer": "",
            "justificacion": "",
            "citas": [""],
            "fuente": "",
            "pagina": null
          }}
        ]
      }}
    ]
  }},
  "alcance": {{
    "comentarios": "",
    "partidas": [
      {{
        "numero": 1,
        "titulo": "",
        "descripcion": "",
        "unidad": "",
        "cantidad": 0,
        "estado": "MATCH",
        "observaciones": "",
        "fuente": "",
        "pagina": null,
        "cita": ""
      }}
    ]
  }},
  "observaciones": {{
    "secciones": [
      {{
        "title": "",
        "items": [
          {{
            "title": "",
            "text": "",
            "quote": "",
            "source": "",
            "page": null
          }}
        ]
      }}
    ]
  }},
  "partidas": [
    {{"numero": 1, "descripcion": "", "unidad": "", "cantidad": 0, "fuente": "", "pagina": null, "cita": ""}}
  ],
  "checklist_sugerido": [
    {{
      "requisito": "",
      "descripcion": "",
      "criterio_cumplimiento": "",
      "formato": "No aplica",
      "categoria": "Legal-Administrativo",
      "aplicabilidad": "Único",
      "obligatorio": "Sí",
      "cumplimiento": "-",
      "status": "Pendiente",
      "prioridad": "Media",
      "fuente": "",
      "pagina": null,
      "cita": ""
    }}
  ],
  "citas": {{}}
}}

REGLAS CRÍTICAS:
- Devuelve SOLO JSON válido, sin markdown.
- Usa únicamente el contenido situado después de cada encabezado "--- DOCUMENTO: nombre ---".
- No inventes, no completes con conocimiento general y no reutilices datos de otras licitaciones.
- Cuando un dato no esté explícito usa exactamente "No se encontró información".
- El nombre de fuente debe coincidir EXACTAMENTE con el encabezado del documento.
- La página debe tomarse de la marca [PAGINA X] anterior a la evidencia.
- Toda respuesta distinta de "No se encontró información" debe tener cita literal, fuente y página.
- Los riesgos permitidos son: ALTO, MEDIO, BAJO y NULO.
- En matriz.secciones analiza requisitos generales, legales, administrativos, técnicos, contractuales, garantías, sanciones, evaluación y causas de desechamiento.
- En financiero.secciones analiza capacidad financiera, pago, facturación, anticipos, retenciones, fianzas, garantías, penas y deducciones.
- En eventos extrae únicamente vigencias y plazos reales.
- En observaciones detecta solo contradicciones, ambigüedades, requisitos restrictivos, posibles candados, fechas incompatibles o condiciones inusuales sustentadas textualmente.
- No generes observaciones si no existe evidencia suficiente.
- En alcance.partidas incluye todos los bienes o servicios solicitados. No inventes cantidades.
- Genera entre 20 y 50 requisitos de checklist cuando el documento lo permita.
- resumen_ejecutivo debe contener exactamente las 16 preguntas indicadas y en el mismo orden.
- Para cada valor de ficha, fechas_clave y resumen_ejecutivo crea una entrada en citas con clave como "ficha.numero_licitacion", "fechas_clave.fallo" o "resumen_ejecutivo.0".
- Cada cita literal debe tener máximo 350 caracteres.

TEXTO DE LOS DOCUMENTOS:
{compact}
""".strip()

def _get_openai_models() -> List[str]:
    """
    Forzado a gpt-5.4-mini para este servicio.
    No usa gpt-5.5, no usa fallback y no lee otro modelo aunque exista en .env.
    """
    return ["gpt-5.4-mini"]


def _extract_text_from_response_payload(payload: Dict[str, Any]) -> str:
    """
    Extrae texto desde Responses API sin depender del SDK de OpenAI.
    """
    if isinstance(payload.get("output_text"), str) and payload["output_text"].strip():
        return payload["output_text"].strip()

    parts = []

    for item in payload.get("output", []) or []:
        for block in item.get("content", []) or []:
            text = block.get("text")

            if isinstance(text, str) and text.strip():
                parts.append(text.strip())

    return "\n".join(parts).strip()


def _parse_model_json(model: str, content: str) -> Dict[str, Any]:
    clean = (content or "").strip()

    if clean.startswith("```"):
        clean = re.sub(r"^```(?:json)?\s*", "", clean, flags=re.IGNORECASE)
        clean = re.sub(r"\s*```$", "", clean)

    try:
        return json.loads(clean)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", clean, re.DOTALL)

        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                pass

        raise Exception(f"Modelo {model} devolvio JSON invalido: {content[:500]}")


def _call_openai_once(model: str, raw_text: str) -> Dict[str, Any]:
    if not OPENAI_API_KEY:
        raise Exception("Falta OPENAI_API_KEY en .env")

    base_url = os.getenv("OPENAI_BASE_URL", "https://api.openai.com").rstrip("/")
    url = f"{base_url}/v1/responses"

    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json",
    }

    timeout_seconds = int(os.getenv("OPENAI_TIMEOUT", "300"))

    body = {
        "model": model,
        "input": [
            {
                "role": "system",
                "content": LICITACION_SYSTEM_PROMPT,
            },
            {
                "role": "user",
                "content": _build_licitacion_prompt(raw_text),
            },
        ],
        "text": {
            "format": {
                "type": "json_object"
            }
        },
    }

    response = requests.post(
        url,
        headers=headers,
        json=body,
        timeout=timeout_seconds,
    )

    if not response.ok:
        raise Exception(
            f"OpenAI Responses API error {response.status_code}: {response.text[:1000]}"
        )

    payload = response.json()
    content = _extract_text_from_response_payload(payload)

    if not content:
        raise Exception(
            f"Modelo {model} no devolvio contenido. Payload: "
            f"{json.dumps(payload, ensure_ascii=False)[:1000]}"
        )

    return _parse_model_json(model, content)


def structure_licitacion_with_openai(raw_text: str) -> Dict[str, Any]:
    models = _get_openai_models()

    if not models:
        raise Exception("No hay modelos OpenAI configurados")

    last_error: Optional[Exception] = None

    for model in models:
        try:
            _log(f"OpenAI: probando modelo {model}")
            return _call_openai_once(model, raw_text)
        except Exception as e:
            msg = str(e)
            _log(f"OpenAI fallo con modelo {model}: {msg}")
            last_error = e

    raise last_error or Exception("Todos los modelos OpenAI fallaron")


# ════════════════════════════════════════════════════════════
# PIPELINE PRINCIPAL
# ════════════════════════════════════════════════════════════

def process_files(file_paths: List[str], include_raw: bool = False) -> Dict[str, Any]:
    if not file_paths:
        raise Exception("No se recibieron archivos")

    combined_parts = []
    documents_info = []

    for fp in file_paths:
        p = Path(fp)
        name = p.name

        if not p.exists():
            documents_info.append({
                "file": name,
                "status": "error",
                "error": "Archivo no existe",
            })
            continue

        ext = p.suffix.lower()

        try:
            _log(f"Procesando {name}...")

            if ext == ".pdf":
                with open(p, "rb") as f:
                    file_bytes = f.read()

                result = analyze_pdf_with_auto_split(file_bytes)
                analyze_result = result.get("analyzeResult", {}) or {}
                content = azure_result_to_page_marked_text(analyze_result)

            elif ext in (".docx", ".doc"):
                content = "[PAGINA desconocida]\n" + extract_docx_text(str(p))

            else:
                documents_info.append({
                    "file": name,
                    "status": "error",
                    "error": f"Extension no soportada: {ext}",
                })
                continue

            if content:
                combined_parts.append(f"--- DOCUMENTO: {name} ---\n{content}")
                page_numbers = [
                    int(n) for n in re.findall(r"\[PAGINA\s+(\d+)\]", content, flags=re.IGNORECASE)
                ]
                documents_info.append({
                    "file": name,
                    "status": "ok",
                    "text_length": len(content),
                    "extracted_text": content,
                    "extracted_raw": {
                        "pages_count": max(page_numbers) if page_numbers else None,
                        "pages": page_numbers,
                        "summary": _clean_extracted_value(content, 1200),
                    },
                    "raw_preview": _clean_extracted_value(content, 1200),
                })
            else:
                documents_info.append({
                    "file": name,
                    "status": "empty",
                    "error": "Sin contenido extraido",
                })

        except Exception as e:
            _log(f"Error con {name}: {e}")
            documents_info.append({
                "file": name,
                "status": "error",
                "error": str(e),
            })

    combined_text = "\n\n".join(combined_parts)

    if not combined_text:
        return {
            "ok": False,
            "error": "No se pudo extraer texto de ningun archivo",
            "documents": documents_info,
        }

    try:
        _log("Estructurando con OpenAI...")
        structured = structure_licitacion_with_openai(combined_text)

        _log("Postprocesando citas, fuentes y paginas...")
        structured = postprocess_structured_data(structured, combined_text)

    except Exception as e:
        return {
            "ok": False,
            "error": f"Error OpenAI: {e}",
            "documents": documents_info,
        }

    output = {
        "ok": True,
        "structured": structured,
        "documents": documents_info,
    }

    output["raw_text_combined"] = combined_text

    return output


def main():
    parser = argparse.ArgumentParser(
        description="Extrae y estructura licitaciones con Azure Document Intelligence + OpenAI"
    )

    parser.add_argument(
        "files",
        nargs="+",
        help="Rutas absolutas a archivos PDF/DOCX",
    )

    parser.add_argument(
        "--include-raw",
        action="store_true",
        help="Incluir texto crudo combinado en la salida",
    )

    args = parser.parse_args()

    try:
        result = process_files(args.files, include_raw=args.include_raw)
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(0 if result.get("ok") else 1)

    except Exception as e:
        print(json.dumps({
            "ok": False,
            "error": str(e),
        }, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()